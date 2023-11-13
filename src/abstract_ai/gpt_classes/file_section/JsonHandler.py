import os
import re
import json
from abstract_utilities.json_utils import get_value_from_path,get_any_value,safe_json_loads,get_any_value,safe_read_from_json
from abstract_utilities import read_from_file
def read_saved_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data
def read_docx(file_path):
    # Load the document
    doc = Document(file_path)
    
    # Read and print each paragraph in the document
    text = ''
    for paragraph in doc.paragraphs:
        text+='\n'+paragraph
    return eatAll(text,['\n','',' ','\t'])
def replace_place_holder_ends(string,placeholder):
    # Find the position of the placeholder
    index = string.find(placeholder)
    if index != -1:
        # Replace from the start of the placeholder to the end of the string
        string = '"'+string[index + len(placeholder):]
    # Find the position of the placeholder from the end
    index = string.rfind(placeholder)
    if index != -1:
        # Replace from the beginning of the string to the start of the placeholder
        string = string[:index]+'"'
    return string
def process_all_strins(all_temp_string,strings,all_js):
    strings.append(replace_place_holder_ends(all_temp_string,'<<**D'))
    if len(strings)==2:
        all_js[strings[0][1:-1]]=strings[1].replace('<<**D',"\'")[1:-1]
        strings = []
    all_temp_string=''
    return all_temp_string,strings,all_js
def fix_malformed_json(malformed_json_string):
    temp_string = ''
    pre_temp_string=''
    all_temp_string=''
    all_js={}
    strings=[]
    for char in malformed_json_string:
        if char in [',', ':']:
            if enclosure[0] == ['"',"'"][0]:
                replace = ['"',"'"][1]
            else:
                replace = ['"',"'"][0]
            all_temp_string,strings,all_js=process_all_strins(all_temp_string,strings,all_js)
            enclosure=[]
        elif char in ['"',"'"]:
            if '<<**D' not in all_temp_string:
                enclosure=[char,char]
            all_temp_string+='<<**D'
            enclosure[-1]=char
        else:
            all_temp_string+=char
    
    # Add any remaining temp_string content
    all_temp_string,strings,all_js=process_all_strins(all_temp_string,strings,all_js)
    return all_js
def preprocess_string(s):
    # Replace \" with a unique marker
    s = s.replace('\\"', '*<d<*')
    # Replace any remaining " with '
    s = s.replace('"', "'")
    # Replace the original single quotes (now all are delimiters) with "
    s = s.replace("'", '"')
    # Replace the unique markers back to \"
    s = s.replace('*<d<*', '\\"')
    return s
def postprocess_string(s):
    # After loading the JSON and processing it, if you need to output the JSON
    # or parts of it as a string with the original quote characters, do the reverse.
    s = s.replace('\\"', '*<d<*')
    s = s.replace('"', "'")
    s = s.replace("*<d<*", '\\"')
    input(s)
    return s
def replace(data):
    return re.sub(r"(\s*'\s*|\s*'\s*:\s*'\s*|\s*'\s*,\s*'\s*)", lambda x: x.group().replace("'", '*<s<*'), data)

def get_any_value_converter(output,key,return_dict=False,replacements=None):
    return_desired_dict=False
    # Preprocess the string to temporarily replace quotes
    if isinstance(output,dict):
        if return_dict:
            return output
        return output[key]
    json_ready_string = preprocess_string(output)
    
    try:
        # Load the string as a JSON object
        data_json = json.loads(json_ready_string)
        value = get_any_value(data_json,key)
        # Example of processing and then converting back to a string
        output_string = json.dumps(value, indent=4)
        input(output_string)
        output_string=postprocess_string(output_string.replace(replacements,"'"))
        return_desired_dict=data_json
    except json.JSONDecodeError as e:
        # Handle JSON errors
        print("Decoding JSON has failed:", e)
        output_string = fix_malformed_json(output)
    if return_dict and return_desired_dict:
        return return_desired_dict
    return output_string
def find_paths_to_key_in_string(data, key, current_path=None, paths=None):
    if paths is None:
        paths = []
    if current_path is None:
        current_path = []
    if isinstance(data, dict):
        for k, v in data.items():
            new_path = current_path + [k]
            if k == key:
                paths.append(new_path)
            if isinstance(v, (dict, list, str)):
                find_paths_to_key_in_string(v, key, new_path, paths)
    elif isinstance(data, list):
        for i, item in enumerate(data):
            new_path = current_path + [i]
            if isinstance(item, (dict, list, str)):
                find_paths_to_key_in_string(item, key, new_path, paths)
    elif isinstance(data, str):
        # Preprocess string to replace problematic quotes
        processed_string = preprocess_string(data)
        try:
            # Attempt to convert the string to JSON and search within it
            json_data = json.loads(processed_string)
            find_paths_to_key_in_string(json_data, key, current_path, paths)
        except json.JSONDecodeError:
            # If decoding fails, check if the string itself contains the key
            if key in data:
                paths.append(current_path)
    return paths
def get_value_from_object(data=None,key_to_find=None):
    dict_values=[]
    if isinstance(data,str):
        if os.path.exists(data):
            data=read_saved_json(file_path=data)
    data=safe_json_loads(data)
    value = get_any_value(data,key_to_find)
    if isinstance(data,str):
        if key_to_find in data:
            data={"find":data}
    paths = find_paths_to_key_in_string(data,key_to_find)
    for path in paths:
        string_value = get_value_from_path(data,path[:-1])
        dict_values.append(get_any_value_converter(string_value,key_to_find,return_dict=True))
    return dict_values
class FileCollator:
    def __init__(self,files_list,key_value=None):
        self.files_list=files_list or []
        self.key_value=key_value or []
    def get_gollated_responses(self,files_list=None,key_value=None):
        if files_list == None:
            files_list=self.files_list
        if key_value == None:
            key_value=self.key_value
        files = self.get_json_data(files_list,key_value=key_value)
        collated_string = self.collate_responses(files)
        return collated_string
    def collate_responses(self,files_list):
        collate_str=''
        nix_list=[]
        for each in files_list:
            lowest = self.get_oldest_first(files_list,nix_list=nix_list)
            nix_list.append(lowest[0])
            collate_str +='\n'+str(files_list[lowest[0]]["value"])
        return collate_str
    @staticmethod
    def get_json_data(files_list,key_value=None):
        if key_value == None:
            if hasattr(self,'key_value'):
                key_value=self.key_value
            key_value = key_value or 'query_response'
        files = []
        for file_path in files_list:
            data = safe_read_from_json(file_path)
            api_response = get_any_value(get_any_value(data,'response'),'api_response')
            response = get_any_value(data,'response')
            created=get_any_value(response ,'created')
            if isinstance(created,list):
                if len(created)>0:
                    created=created[0]
            files.append({'created':int(created),"value":api_response})
        return files
    @staticmethod
    def get_oldest_first(json_list,nix_list=[]):
        lowest=[None,None]
        for i,values in enumerate(json_list):
            if i not in nix_list:
                if lowest[0] == None:
                    lowest=[i,int(values['created'])]
                elif int(values['created']) < int(lowest[1]):
                    lowest=[i,int(values['created'])]
        return lowest
def read_from_file_with_multiple_encodings(file_path, encodings=None):
    COMMON_ENCODINGS = [
    'utf-8', 
    'utf-16', 
    'utf-16-be', 
    'utf-16-le', 
    'utf-32', 
    'utf-32-be', 
    'utf-32-le',
    'ISO-8859-1', # also known as latin1
    'ISO-8859-2', # Central and Eastern European languages 
    'ISO-8859-3', 
    'ISO-8859-4',
    'ISO-8859-5', # Cyrillic alphabet
    'ISO-8859-6', # Arabic
    'ISO-8859-7', # Greek
    'ISO-8859-8', # Hebrew
    'ISO-8859-9', # Turkish
    'ISO-8859-10',
    'ISO-8859-13',
    'ISO-8859-14',
    'ISO-8859-15',
    'ISO-8859-16',
    'windows-1250',
    'windows-1251',
    'windows-1252',
    'windows-1253',
    'windows-1254',
    'windows-1255',
    'windows-1256',
    'windows-1257',
    'windows-1258',
    'big5',
    'big5hkscs',
    'cp037',
    'cp424',
    'cp437',
    'cp500',
    'cp720',
    'cp737',
    'cp775',
    'cp850',
    'cp852',
    'cp855',
    'cp856',
    'cp857',
    'cp858',
    'cp860',
    'cp861',
    'cp862',
    'cp863',
    'cp864',
    'cp865',
    'cp866',
    'cp869',
    'cp874',
    'cp875',
    'cp932',
    'cp949',
    'cp950',
    'cp1006',
    'cp1026',
    'cp1140',
    'cp1256',
    'euc_jp',
    'euc_jis_2004',
    'euc_jisx0213',
    'euc_kr',
    'gb2312',
    'gbk',
    'gb18030',
    'hz',
    'iso2022_jp',
    'iso2022_jp_1',
    'iso2022_jp_2',
    'iso2022_jp_2004',
    'iso2022_jp_3',
    'iso2022_jp_ext',
    'iso2022_kr',
    'latin_1',
    'koi8_r',
    'koi8_t',
    'koi8_u',
    'mac_cyrillic',
    'mac_greek',
    'mac_iceland',
    'mac_latin2',
    'mac_roman',
    'mac_turkish',
    'ptcp154',
    'shift_jis',
    'shift_jis_2004',
    'shift_jisx0213',
    'utf_32_be',
    'utf_32_le',
    'utf_16_be',
    'utf_16_le',
    'utf_7',
    'utf_8_sig',
    'latin-1']
    if encodings is None:
        encodings = COMMON_ENCODINGS

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue  # Try the next encoding if decoding fails

    # If none of the encodings work, return None or handle the error as needed
    return None
from docx import Document
def read_docx(file_path):
    # Load the document
    doc = Document(file_path)
    
    # Read and print each paragraph in the document
    text = ''
    for paragraph in doc.paragraphs:
        text+='\n'+paragraph
    return eatAll(text,['\n','',' ','\t'])


